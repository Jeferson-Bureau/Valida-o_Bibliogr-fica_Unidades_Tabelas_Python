import sys
import csv
import os
import argparse
from typing import List, Tuple
from rich.console import Console
from rich.table import Table
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models import ReferenceInput, ValidationResult
from validators import CrossrefValidator, OpenAlexValidator, PubMedValidator, SemanticScholarValidator
from word_parser import extract_references_from_docx
from reference_parser import parse_abnt_reference
from apa7_formatter import format_apa7
from table_formatter import format_all_tables_in_docx

console = Console()

# ──────────────────────────────────────────────
# Mapeamento de cores por status
# ──────────────────────────────────────────────
STATUS_CONFIG = {
    "Confirmado": {
        "emoji": "✅",
        "color": RGBColor(0x1E, 0x8B, 0x1E),   # verde escuro
        "label": "Autoria e Veracidade CONFIRMADAS",
    },
    "Parcialmente confirmado": {
        "emoji": "⚠️",
        "color": RGBColor(0xCC, 0x66, 0x00),   # laranja escuro
        "label": "PARCIALMENTE confirmada",
    },
    "Não confirmado": {
        "emoji": "❓",
        "color": RGBColor(0xCC, 0x00, 0x00),   # vermelho
        "label": "NÃO CONFIRMADA — Referência não localizada nas bases consultadas",
    },
    "Possível referência falsa": {
        "emoji": "🚫",
        "color": RGBColor(0x99, 0x00, 0x00),   # vermelho intenso
        "label": "POSSÍVEL REFERÊNCIA FALSA",
    },
    "Erro": {
        "emoji": "⚙️",
        "color": RGBColor(0x88, 0x88, 0x88),   # cinza
        "label": "ERRO na consulta às APIs",
    },
}

def get_status_config(status: str) -> dict:
    status_lower = status.lower()
    for key in STATUS_CONFIG:
        if key.lower() == status_lower:
            return STATUS_CONFIG[key]
    return STATUS_CONFIG["Não confirmado"]


# ──────────────────────────────────────────────
# Classe principal de validação
# ──────────────────────────────────────────────
class ReferenceValidator:
    def __init__(self):
        self.validators = [
            CrossrefValidator(),
            OpenAlexValidator(),
            PubMedValidator(),
            SemanticScholarValidator()
        ]

    def validate_reference(self, reference: ReferenceInput) -> ValidationResult:
        best_result = None
        for validator in self.validators:
            result = validator.validate(reference)
            if result and result.status != "Erro":
                if result.status == "Confirmado":
                    return result
                if best_result is None or (
                    result.status == "Parcialmente confirmado"
                    and best_result.status != "Parcialmente confirmado"
                ):
                    best_result = result
            elif result and result.status == "Erro":
                if best_result is None:
                    best_result = result

        if best_result:
            return best_result

        return ValidationResult(
            api_source="Nenhuma",
            status="Não confirmado",
            issues=["Referência não encontrada em nenhuma das bases consultadas."],
            confidence="Baixo"
        )


# ──────────────────────────────────────────────
# Inserção de anotação colorida no parágrafo do Word
# ──────────────────────────────────────────────
def annotate_paragraph(para, result: ValidationResult):
    """
    Adiciona um Run colorido ao final do parágrafo da referência com o resultado da validação.
    """
    cfg = get_status_config(result.status)

    # Linha 1: selo de status
    run = para.add_run(f"  [{cfg['label']}]")
    run.bold = True
    run.font.color.rgb = cfg["color"]
    run.font.size = Pt(9)

    # Linha 2: detalhes (autores encontrados, fonte, DOI, observações)
    details_parts = []

    if result.authors_found:
        details_parts.append(f"Autores encontrados: {', '.join(result.authors_found[:5])}")

    if result.doi_found:
        details_parts.append(f"DOI: {result.doi_found}")

    if result.issues:
        details_parts.append(f"Divergencias: {' | '.join(result.issues)}")

    details_parts.append(f"Fonte: {result.api_source} | Confianca: {result.confidence}")

    if details_parts:
        run2 = para.add_run("\n    " + " · ".join(details_parts))
        run2.italic = True
        run2.font.color.rgb = cfg["color"]
        run2.font.size = Pt(8)

    # Linha 3: referência APA7 sugerida
    apa7 = format_apa7(result)
    if apa7:
        run3 = para.add_run(f"\n    APA7: {apa7}")
        run3.italic = True
        run3.font.color.rgb = RGBColor(0x33, 0x66, 0x99)  # azul discreto
        run3.font.size = Pt(8)


# ──────────────────────────────────────────────
# Exibição de relatório no terminal
# ──────────────────────────────────────────────
def print_report(results: List[Tuple[ReferenceInput, ValidationResult]]):
    table = Table(title="Relatório de Validação de Referências", show_lines=True)

    table.add_column("Nº", justify="right", style="cyan", no_wrap=True, width=4)
    table.add_column("Referência Original", style="magenta", max_width=40)
    table.add_column("Status", style="bold", width=22)
    table.add_column("Autoria", width=10)
    table.add_column("Fonte", style="blue", width=14)
    table.add_column("DOI / Link", style="cyan", max_width=30)
    table.add_column("Observações", style="red", max_width=35)

    for idx, (ref, res) in enumerate(results, 1):
        if res.status == "Confirmado":
            status_color = "green"
            autoria = "[green]Sim[/green]"
        elif "Parcial" in res.status:
            status_color = "yellow"
            autoria = "[yellow]Parcial[/yellow]"
        else:
            status_color = "red"
            autoria = "[red]Nao[/red]"

        status_text = f"[{status_color}]{res.status}[/{status_color}]"
        issues_text = " | ".join(res.issues) if res.issues else "—"
        link_text = res.doi_found or (res.url[:40] if res.url else "—")

        table.add_row(
            str(idx),
            ref.original_text[:80] + "…" if len(ref.original_text) > 80 else ref.original_text,
            status_text,
            autoria,
            res.api_source,
            link_text,
            issues_text
        )

    console.print(table)


# ──────────────────────────────────────────────
# Ponto de entrada
# ──────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(
        description="Validador de Referências Bibliográficas — insere o resultado direto no Word"
    )
    parser.add_argument("file", nargs="?", help="Caminho para o arquivo .docx")
    parser.add_argument(
        "--output", "-o",
        default=None,
        help="Caminho para o arquivo Word de saída (padrão: <nome_original>_verificado.docx)"
    )
    parser.add_argument(
        "--csv",
        default="relatorio_referencias.csv",
        help="Caminho para o relatório CSV (padrão: relatorio_referencias.csv)"
    )
    parser.add_argument(
        "--format-tables",
        action="store_true",
        default=True,
        help="Aplica formatação padrão Acta Scientiarum/APA nas tabelas do documento (padrão: Ativo)"
    )
    parser.add_argument(
        "--no-format-tables",
        dest="format_tables",
        action="store_false",
        help="Desativa a formatação automática de tabelas"
    )
    args = parser.parse_args()

    validator = ReferenceValidator()

    if args.file:
        console.print(f"\n[bold blue]Lendo: {args.file}[/bold blue]")
        try:
            doc, ref_paragraphs = extract_references_from_docx(args.file)
        except Exception as e:
            console.print(f"[bold red]Erro ao ler o arquivo Word: {e}[/bold red]")
            sys.exit(1)

        if not ref_paragraphs:
            console.print("[yellow]Nenhuma referencia encontrada. Verifique se a secao se chama 'Referencias'.[/yellow]")
            sys.exit(0)

        console.print(f"[green]{len(ref_paragraphs)} referencia(s) encontrada(s). Iniciando validacao...[/green]\n")

        references = []
        for para_obj, ref_text in ref_paragraphs:
            parsed = parse_abnt_reference(ref_text)
            references.append((
                para_obj,
                ReferenceInput(
                    id=ref_text[:40],
                    original_text=ref_text,
                    title=parsed.get("title"),
                    doi=parsed.get("doi"),
                    authors=parsed.get("authors", [])
                )
            ))

    else:
        console.print("[yellow]Nenhum arquivo fornecido. Rodando exemplos de teste internos.[/yellow]\n")
        doc = None
        references = [
            (None, ReferenceInput(
                id="1",
                original_text="Vaswani, Ashish, et al. Attention is all you need. Advances in neural information processing systems 30 (2017).",
                title="Attention is all you need",
                authors=["Vaswani, Ashish", "Shazeer, Noam"]
            )),
            (None, ReferenceInput(
                id="2",
                original_text="Artigo inexistente para testar erro. Silva, J.",
                title="Análise de algoritmos super obscuros de marte",
                authors=["Silva, J."]
            ))
        ]

    # ── Validação e anotação ──
    results = []
    with console.status("[bold green]Consultando bases academicas...") as _:
        for para_obj, ref in references:
            console.log(f"  Validando: {ref.original_text[:70]}...")
            res = validator.validate_reference(ref)
            results.append((ref, res))

            # Anota diretamente no parágrafo correto da lista de referências
            if doc is not None and para_obj is not None:
                annotate_paragraph(para_obj, res)

    # ── Relatório no terminal ──
    console.print()
    print_report(results)

    # ── Salva o Word anotado ──
    if doc is not None and args.file:
        if args.format_tables and len(doc.tables) > 0:
            count = format_all_tables_in_docx(doc)
            console.print(f"[bold cyan]{count} tabela(s) formatada(s) no padrão Acta Scientiarum (3 linhas).[/bold cyan]")

        output_path = args.output
        if not output_path:
            base, ext = os.path.splitext(args.file)
            output_path = f"{base}_verificado{ext}"

        doc.save(output_path)
        console.print(f"\n[bold green]Documento Word verificado salvo em:[/bold green] {output_path}")

    # ── Salva relatório CSV ──
    try:
        with open(args.csv, "w", encoding="utf-8-sig", newline="") as f:
            writer = csv.writer(f)
            writer.writerow(["Nº", "Referência Original", "Status", "Autoria Confirmada",
                             "DOI/Link", "Autores Encontrados", "Fonte", "Confiança",
                             "Observações", "Referência APA7"])
            for idx, (ref, res) in enumerate(results, 1):
                autoria = "Sim" if res.status == "Confirmado" else ("Parcial" if "Parcial" in res.status else "Não")
                apa7 = format_apa7(res) or "—"
                writer.writerow([
                    idx,
                    ref.original_text,
                    res.status,
                    autoria,
                    res.doi_found or res.url or "—",
                    "; ".join(res.authors_found),
                    res.api_source,
                    res.confidence,
                    " | ".join(res.issues) if res.issues else "—",
                    apa7,
                ])
        console.print(f"[bold green]Relatorio CSV salvo em:[/bold green] {args.csv}")
    except Exception as e:
        console.print(f"[bold red]Erro ao salvar CSV: {e}[/bold red]")


if __name__ == "__main__":
    main()
