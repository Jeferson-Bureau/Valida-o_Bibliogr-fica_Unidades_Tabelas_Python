"""
si_units_corrector.py — Corretor de Unidades Científicas para Python/python-docx
Porta da Macro VBA "Sistema Internacional de Unidades" v2.1
Normas: APA 7ª Edição + Sistema Internacional de Unidades (SI / BIPM)
"""

import re
from typing import List
import docx
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.text.run import Run

# ──────────────────────────────────────────────────────────────────────────────
# Caracteres especiais (mesmos do VBA)
# ──────────────────────────────────────────────────────────────────────────────
DEG   = "\u00B0"       # °  símbolo de grau correto
ORD   = "\u00BA"       # º  ordinal (errado para grau)
MU    = "\u03BC"       # μ  micro
OHM   = "\u03A9"       # Ω  Ohm
SM1   = "\u207B\u00B9" # ⁻¹
SM2   = "\u207B\u00B2" # ⁻²
SM3   = "\u207B\u00B3" # ⁻³
SUP2  = "\u00B2"       # ²
SUP3  = "\u00B3"       # ³
SUPn  = "\u207F"       # ⁿ
SUB0  = "\u2080"       # ₀
SUPf  = "\u1DA0"       # ᶠ
NBSP  = "\u00A0"       # espaço não quebrável
PH_DEG = "\uE000"      # placeholder PUA para °C durante bloco F
PRIME  = "\u2032"      # ′ minuto
DPRIME = "\u2033"      # ″ segundo

# ──────────────────────────────────────────────────────────────────────────────
# Compilação de Expressões Regulares e Mapas
# ──────────────────────────────────────────────────────────────────────────────

# Bloco A - Micro (ASCII u -> μ)
PAIRS_MICRO_RE = [
    (re.compile(r'(\d)\s?um\b'),   rf'\1 {MU}m'),
    (re.compile(r'(\d)\s?ug\b'),   rf'\1 {MU}g'),
    (re.compile(r'(\d)\s?uL\b'),   rf'\1 {MU}L'),
    (re.compile(r'(\d)\s?us\b'),   rf'\1 {MU}s'),
    (re.compile(r'(\d)\s?umol\b'), rf'\1 {MU}mol'),
]

# Bloco B - Plurais indevidos
PLURAL_MAP = {
    "kgs":"kg","mgs":"mg","ngs":"ng","kms":"km","cms":"cm","mms":"mm","nms":"nm",
    "mLs":"mL","dLs":"dL","mins":"min",
    "kHzs":"kHz","MHzs":"MHz","GHzs":"GHz",
    "kPas":"kPa","MPas":"MPa","GPas":"GPa",
    "kNs":"kN","MNs":"MN","kJs":"kJ","MJs":"MJ","kWs":"kW","MWs":"MW",
    "mVs":"mV","kVs":"kV","mAs":"mA",
    "mols":"mol","mmols":"mmol","atms":"atm","Torrs":"Torr","mTs":"mT",
}
# Regex única para encontrar as chaves
PLURAL_RE = re.compile(r'\b(' + '|'.join(re.escape(k) for k in PLURAL_MAP.keys()) + r')\b')

# Bloco C - Prefixos soltos (ex: "k m" -> "km")
PREFIX_PAIRS = {
    "k m":"km","k g":"kg","k Hz":"kHz","k Pa":"kPa","k N":"kN",
    "k J":"kJ","k W":"kW","k V":"kV",
    "M Hz":"MHz","M Pa":"MPa","M N":"MN","M J":"MJ","M W":"MW","M V":"MV",
    "G Hz":"GHz","G Pa":"GPa",
    "m m":"mm","m L":"mL","m V":"mV","m A":"mA","m T":"mT",
    "m mol":"mmol","m g":"mg",
    "n m":"nm","n g":"ng","n s":"ns",
    f"{MU} m":f"{MU}m",f"{MU} g":f"{MU}g",f"{MU} L":f"{MU}L",
    f"{MU} s":f"{MU}s",f"{MU} mol":f"{MU}mol",
}
PREFIX_RE = re.compile(r'(\d)\s?(' + '|'.join(re.escape(k) for k in PREFIX_PAIRS.keys()) + r')')
def _prefix_sub(match):
    return f"{match.group(1)} {PREFIX_PAIRS[match.group(2)]}"

# Bloco G - Unidades Compostas
COMPOUND_MAP_LIST = [
    # Massa/volume
    (f"{MU}g/mL",  f"{MU}g mL{SM1}"), (f"{MU}g/dL",  f"{MU}g dL{SM1}"),
    (f"{MU}g/L",   f"{MU}g L{SM1}"),   (f"{MU}g/g",   f"{MU}g g{SM1}"),
    (f"{MU}g/100 g", f"{MU}g 100g{SM1}"), (f"{MU}g/100g", f"{MU}g 100g{SM1}"),
    ("ng/mL",  f"ng mL{SM1}"),  ("ng/L",  f"ng L{SM1}"),
    ("mg/mL",  f"mg mL{SM1}"),  ("mg/dL", f"mg dL{SM1}"),
    ("mg/ml",  f"mg mL{SM1}"),  ("mg/dl", f"mg dL{SM1}"),
    ("mg/L",   f"mg L{SM1}"),   ("mg/kg", f"mg kg{SM1}"),
    ("U/mL",   f"U mL{SM1}"),   ("U/ml",  f"U mL{SM1}"),
    ("U/mg",   f"U mg{SM1}"),
    ("10n/q",  f"10{SUPn} q{SM1}"),
    ("mg/100 g",f"mg 100g{SM1}"), ("mg/100g",f"mg 100g{SM1}"),
    ("g/mL",   f"g mL{SM1}"),   ("g/dL",  f"g dL{SM1}"),
    ("g/L",    f"g L{SM1}"),    ("g/kg",  f"g kg{SM1}"),
    
    # Substância/volume
    (f"{MU}mol/L",  f"{MU}mol L{SM1}"),
    ("nmol/L", f"nmol L{SM1}"),
    ("mmol/mL",f"mmol mL{SM1}"),("mmol/L", f"mmol L{SM1}"),
    ("mol/mL", f"mol mL{SM1}"), ("mol/L",  f"mol L{SM1}"),
    (f"mol/m{SUP3}",f"mol m{SM3}"), ("mol/m3", f"mol m{SM3}"),
    
    # Velocidade / aceleração
    (f"m/s{SUP2}", f"m s{SM2}"), ("m/s2", f"m s{SM2}"),
    ("km/h",  f"km h{SM1}"),  ("km/s", f"km s{SM1}"),
    ("m/s",   f"m s{SM1}"),
    
    # Força / pressão / energia / potência
    (f"N/m{SUP2}", f"N m{SM2}"), ("N/m2", f"N m{SM2}"),
    (f"W/m{SUP2}", f"W m{SM2}"), ("W/m2", f"W m{SM2}"),
    ("N/m",   f"N m{SM1}"),   ("Pa/m",  f"Pa m{SM1}"),
    ("kJ/mol/nm", f"kJ mol{SM1} nm{SM1}"),
    ("J/mol", f"J mol{SM1}"), ("J/kg",  f"J kg{SM1}"),
    ("J/g",   f"J g{SM1}"),   ("W/kg",  f"W kg{SM1}"),
    ("Kcal/g", f"Kcal g{SM1}"),
    
    # Densidade / Massa por área
    (f"kg/m{SUP3}", f"kg m{SM3}"), ("kg/m3", f"kg m{SM3}"),
    (f"kg/m{SUP2}", f"kg m{SM2}"), ("kg/m2", f"kg m{SM2}"),
    (f"g/cm{SUP3}", f"g cm{SM3}"), ("g/cm3", f"g cm{SM3}"),
    
    # Fluxo / débito
    ("mL/min",f"mL min{SM1}"),("L/min", f"L min{SM1}"),
    ("mL/h",  f"mL h{SM1}"),  ("L/h",   f"L h{SM1}"),
    ("mL/kg", f"mL kg{SM1}"), ("L/kg",  f"L kg{SM1}"),
    ("min/day",f"min day{SM1}"),("min/dia",f"min dia{SM1}"),
    ("min/d", f"min d{SM1}"),
    
    # Compostos / antioxidantes
    ("GAE/kg",  f"GAE kg{SM1}"),  ("GAE/g",    f"GAE g{SM1}"),
    ("GAE/100 g",f"GAE 100g{SM1}"),("GAE/100g",f"GAE 100g{SM1}"),
    ("QE/100 g",f"QE 100g{SM1}"), ("QE/100g",  f"QE 100g{SM1}"),
    ("TEAC/100 g",f"TEAC 100g{SM1}"),("TEAC/100g",f"TEAC 100g{SM1}"),
    ("catechins/100 g",f"catechins 100g{SM1}"),("catechins/100g",f"catechins 100g{SM1}"),
    ("Trolox/kg",f"Trolox kg{SM1}"),
    ("ug/g",  f"{MU}g g{SM1}"),
    ("ug/L",  f"{MU}g L{SM1}"),  ("ug/mL", f"{MU}g mL{SM1}"),
    ("ug/100 g",f"{MU}g 100g{SM1}"),("ug/100g",f"{MU}g 100g{SM1}"),
    
    # Microbiologia: CFU/UFC
    ("CFU/mL",  f"UFC mL{SM1}"),  ("CFU/ml",  f"UFC mL{SM1}"),
    ("CFU/L",   f"UFC L{SM1}"),
    ("CFU/g",   f"UFC g{SM1}"),
    (f"CFU/cm{SUP2}", f"UFC cm{SM2}"), ("CFU/cm2", f"UFC cm{SM2}"),
    (f"CFU/m{SUP2}",  f"UFC m{SM2}"),  ("CFU/m2",  f"UFC m{SM2}"),
    ("CFU/mL/h",f"UFC mL{SM1} h{SM1}"),
    ("UFC/mL",  f"UFC mL{SM1}"),  ("UFC/ml",  f"UFC mL{SM1}"),
    ("UFC/L",   f"UFC L{SM1}"),
    ("UFC/g",   f"UFC g{SM1}"),
    (f"UFC/cm{SUP2}", f"UFC cm{SM2}"), ("UFC/cm2", f"UFC cm{SM2}"),
    (f"UFC/m{SUP2}",  f"UFC m{SM2}"),  ("UFC/m2",  f"UFC m{SM2}"),
    
    # Temperatura / taxa
    (f"{DEG}C/min", f"{DEG}C min{SM1}"),

    # Novas regras solicitadas
    (f"{MU}g/ml", f"{MU}g ml{SM1}"),
    (f"{MU}g h/ml", f"{MU}g h ml{SM1}"),
    ("pg/ml", f"pg ml{SM1}"),
    (f"{MU}mol/g", f"{MU}mol g{SM1}"),
    ("IU/L", f"IU L{SM1}"),
]

# Compila um dicionário e regex otimizada para substituição múltipla (O(1))
# As chaves maiores primeiro para evitar problemas de substring (e.g. kJ/mol vs kJ/mol/nm)
COMPOUND_MAP_LIST.sort(key=lambda x: len(x[0]), reverse=True)
COMPOUND_MAP = {wrong: correct for wrong, correct in COMPOUND_MAP_LIST}
COMPOUND_RE = re.compile('(' + '|'.join(re.escape(k) for k in COMPOUND_MAP.keys()) + ')')

# Lista de unidades SI simples para bloco H
SI_UNITS = {
    "kg","g","mg","ng","pg","t",
    "km","m","cm","mm","nm","pm",
    "kL","L","mL","dL",
    "GHz","MHz","kHz","Hz",
    "GPa","MPa","kPa","Pa",
    "MN","kN","N",
    "MJ","kJ","J","cal","kcal",
    "MW","kW","W",
    "MV","kV","V","mV",
    f"G{OHM}",f"M{OHM}",f"k{OHM}",OHM,"mA","A",
    "T","mT",
    "mmol","mol",
    "Bq","Gy","Sv",
    "bar","atm","mmHg","Torr",
    "h","min","s","ms","ns","ps",
    "K",
    f"{MU}m",f"{MU}g",f"{MU}L",f"{MU}s",f"{MU}mol",
}
# Número seguido de unidade sem espaço
UNIT_RE = re.compile(r'(\d)(' + '|'.join(re.escape(u) for u in sorted(SI_UNITS, key=len, reverse=True)) + r')(?=\s|$|[,;.)\]])')
MU_UNITS_RE = re.compile(r'(\d)' + re.escape(MU) + r'(m|g|L|s|mol)')
CENTRIFUGE_RE = re.compile(r'(\d[\d.,]*)\s*[xX\u00D7]\s*g\b')

# Outras Expressões regulares comuns
OHM_RE = re.compile(r'(?i) ohm\b')
MULTI_SPACE_RE = re.compile(r'  +')
PERCENT_RE = re.compile(r'(\d)\s+%')
DEG_SPACE_RE = re.compile(r'(\d)\s' + re.escape(DEG))
PRIME_SPACE_RE = re.compile(r'(\d)\s' + re.escape(PRIME))
DPRIME_SPACE_RE = re.compile(r'(\d)\s' + re.escape(DPRIME))
CELSIUS_RE = re.compile(r'(\d)\s*' + re.escape(DEG) + r'\s*C\b')
M_M_RE = re.compile(r'\bm/m\b')
P_P0_RE = re.compile(r'\bP\s*/\s*P0\b')
PUNCT_RE = re.compile(r' ([.,;:])')

OPERATORS = ["=", "<", ">", "±", "~", "≤", "≥", "≠", "≈"]
OPERATOR_CHARS = set(OPERATORS)
# Pré-compilação para espaçamento de operadores
OP_SPACING_RE1 = []
OP_SPACING_RE2 = []
for op in OPERATORS:
    e = re.escape(op)
    OP_SPACING_RE1.append((re.compile(rf'(\S)\s*{e}\s*(\S)'), op))
    OP_SPACING_RE2.append((re.compile(rf'(^|[\s\(\[\{{\'\"]){e}\s*(\d)'), op))

def _apply_all(text: str) -> str:
    """Aplica todas as substituições de texto a uma string."""
    
    # Bloco A — Pré-limpeza
    text = text.replace(NBSP, " ")
    text = text.replace("\u00B5", MU)  # Normaliza U+00B5 (Micro Sign) para U+03BC (Greek Mu)
    text = text.replace(f"{ORD}C", f"{DEG}C")
    text = text.replace("\uFF1D", "=").replace("\uFF1C", "<").replace("\uFF1E", ">")
    
    text = OHM_RE.sub(f' {OHM}', text)
    text = text.replace(f"kO", f"k{OHM}").replace(f"MO", f"M{OHM}").replace(f"GO", f"G{OHM}")

    for pat_re, rep in PAIRS_MICRO_RE:
        text = pat_re.sub(rep, text)

    # Bloco B — Plurais indevidos
    text = PLURAL_RE.sub(lambda m: PLURAL_MAP[m.group(1)], text)

    # Bloco C — Prefixos soltos
    text = PREFIX_RE.sub(_prefix_sub, text)

    # Bloco D — Espaços múltiplos
    text = MULTI_SPACE_RE.sub(' ', text)

    # Bloco E — Porcentagem
    text = PERCENT_RE.sub(r'\1%', text)

    # Bloco F — Temperatura e ângulos
    text = text.replace(f"{DEG}C", f"{PH_DEG}C")
    text = DEG_SPACE_RE.sub(lambda m: m.group(1) + DEG, text)
    text = PRIME_SPACE_RE.sub(lambda m: m.group(1) + PRIME, text)
    text = DPRIME_SPACE_RE.sub(lambda m: m.group(1) + DPRIME, text)
    text = text.replace(f"{PH_DEG}C", f"{DEG}C")
    text = CELSIUS_RE.sub(lambda m: m.group(1) + DEG + 'C', text)

    # Bloco G — Unidades compostas (barra → expoente negativo)
    text = COMPOUND_RE.sub(lambda m: COMPOUND_MAP[m.group(1)], text)
    text = M_M_RE.sub(f'm m{SM1}', text)
    text = P_P0_RE.sub(f'P P{SUB0}{SM1}', text)
    text = text.replace(f"Ca/Si{SUPf}", f"Ca Si{SUPf}{SM1}").replace(f"Sn/Si{SUPf}", f"Sn Si{SUPf}{SM1}")
    text = text.replace("Ca/Sif", f"Ca Si{SUPf}{SM1}").replace("Sn/Sif", f"Sn Si{SUPf}{SM1}")

    for dot in ["\u00B7", "\u22C5", "\u2219", "\u2022"]:
        text = text.replace(dot, " ")

    text = MULTI_SPACE_RE.sub(' ', text)

    # Bloco H — Espaço entre número e unidade SI
    text = UNIT_RE.sub(r'\1 \2', text)
    text = MU_UNITS_RE.sub(r'\1 ' + MU + r'\2', text)
    text = CENTRIFUGE_RE.sub(r'\1 ' + "\u00D7" + ' g', text)

    # Bloco I — Operadores matemáticos
    for pat_re, op in OP_SPACING_RE1:
        text = pat_re.sub(lambda m, o=op: f'{m.group(1)} {o} {m.group(2)}', text)
    for pat_re, op in OP_SPACING_RE2:
        text = pat_re.sub(lambda m, o=op: f'{m.group(1)}{o} {m.group(2)}', text)
    
    for op in OPERATORS:
        text = text.replace(f"( {op} ", f"({op} ")
        text = text.replace(f"[ {op} ", f"[{op} ")

    text = MULTI_SPACE_RE.sub(' ', text)

    # Bloco J — Limpeza final de pontuação
    text = PUNCT_RE.sub(r'\1', text)

    return text

def _is_references_heading(para: Paragraph) -> bool:
    """Retorna True se o parágrafo for o título da seção Referências."""
    text = para.text.strip()
    return bool(re.fullmatch(r'refer[eê]ncias?', text, re.IGNORECASE))

def _process_run(run: Run) -> None:
    """Aplica correções ao texto de um único run."""
    original = run.text
    if not original:
        return
    corrected = _apply_all(original)
    if corrected != original:
        run.text = corrected

def _fix_cross_run_operators(para: Paragraph) -> None:
    """
    Garante espaço antes/depois de operadores matemáticos que estão
    em runs separados dos números ao redor.
    """
    runs = para.runs
    for i, run in enumerate(runs):
        txt = run.text
        if not txt:
            continue

        contains_op = any(op in txt for op in OPERATOR_CHARS)
        if not contains_op:
            continue

        if i > 0:
            prev = runs[i - 1]
            if prev.text and not prev.text.endswith(" "):
                prev.text = prev.text + " "

        if i < len(runs) - 1:
            nxt = runs[i + 1]
            if nxt.text and not nxt.text.startswith(" "):
                nxt.text = " " + nxt.text

        run.text = txt.strip()

def _merge_split_compound_units(para: Paragraph) -> None:
    """
    Funde o expoente no run da base antes do processamento (ex: 'kg/m' e '2').
    """
    for i in range(len(para.runs) - 1):
        r1 = para.runs[i]
        r2 = para.runs[i+1]
        if not r1.text or not r2.text:
            continue
            
        if r1.text.endswith(("/m", "/cm", "/s", "/m ", "/cm ", "/s ")):
            match = re.match(r'^(\s*)([23²³])(?:\s|[,.;)]|$)', r2.text)
            if match:
                spaces = match.group(1)
                exp = match.group(2)
                r1.text = r1.text + spaces + exp
                r2.text = r2.text[len(spaces) + 1:]

def _fix_cross_run_spacing(para: Paragraph) -> None:
    """
    Remove espaços indevidos entre números e os símbolos % ou °C que acabaram
    separados em runs diferentes.
    """
    runs = para.runs
    for i in range(len(runs) - 1):
        r1 = runs[i]
        r2 = runs[i+1]
        if not r1.text or not r2.text:
            continue
            
        t2 = r2.text.lstrip()
        if t2.startswith(("%", "°C", "\u00b0C")):
            t1 = r1.text.rstrip()
            if t1 and t1[-1].isdigit():
                r1.text = t1
                r2.text = t2

def _fix_cross_run_compounds(para: Paragraph) -> None:
    """
    Aplica substituições de unidades compostas ao nível do
    parágrafo inteiro, resolvendo padrões que ficaram divididos entre runs.
    """
    runs = para.runs
    if not runs:
        return

    full_text = "".join(r.text or "" for r in runs)
    corrected = full_text

    # Aplica as mesmas substituições de compostos (via Regex)
    corrected = corrected.replace("\u00B5", MU) # Garante que run merges cruzados também estejam normalizados
    corrected = COMPOUND_RE.sub(lambda m: COMPOUND_MAP[m.group(1)], corrected)
    corrected = M_M_RE.sub(f'm m{SM1}', corrected)
    corrected = P_P0_RE.sub(f'P P{SUB0}{SM1}', corrected)
    corrected = corrected.replace(f"Ca/Si{SUPf}", f"Ca Si{SUPf}{SM1}").replace(f"Sn/Si{SUPf}", f"Sn Si{SUPf}{SM1}")
    corrected = corrected.replace("Ca/Sif", f"Ca Si{SUPf}{SM1}").replace("Sn/Sif", f"Sn Si{SUPf}{SM1}")

    if corrected == full_text:
        return

    pos = 0
    for i, run in enumerate(runs):
        orig_len = len(run.text or "")
        if i < len(runs) - 1:
            run.text = corrected[pos:pos + orig_len]
            pos += orig_len
        else:
            run.text = corrected[pos:]

def _process_paragraph(para: Paragraph) -> None:
    """Processa todos os runs de um parágrafo."""
    _merge_split_compound_units(para)
    for run in para.runs:
        _process_run(run)
    _fix_cross_run_compounds(para)
    _fix_cross_run_operators(para)
    _fix_cross_run_spacing(para)
    for run in para.runs:
        if run.text and "  " in run.text:
            run.text = MULTI_SPACE_RE.sub(" ", run.text)

def correct_si_units(doc: Document) -> int:
    """
    Aplica todas as correções de unidades SI ao documento.
    """
    processed = 0
    in_references = False

    for para in doc.paragraphs:
        if _is_references_heading(para):
            in_references = True
        if in_references:
            continue
        _process_paragraph(para)
        processed += 1

    for section in doc.sections:
        for header_footer in [
            section.header, section.footer,
            section.even_page_header, section.even_page_footer,
            section.first_page_header, section.first_page_footer,
        ]:
            try:
                for para in header_footer.paragraphs:
                    _process_paragraph(para)
                    processed += 1
            except Exception:
                pass

    in_references = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _process_paragraph(para)
                    processed += 1

    return processed
