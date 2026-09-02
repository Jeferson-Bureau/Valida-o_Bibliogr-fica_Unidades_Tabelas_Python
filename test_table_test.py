import docx
from table_formatter import format_all_tables_in_docx

def test_table_formatting():
    # Cria um documento temporário com 1 tabela desformatada
    doc = docx.Document()
    doc.add_paragraph("Tabela 1 - Teste de amostragem")
    
    table = doc.add_table(rows=3, cols=3)
    
    # Cabeçalho
    table.cell(0, 0).text = "Tratamento"
    table.cell(0, 1).text = "Réplica 1"
    table.cell(0, 2).text = "Média"
    
    # Dados
    table.cell(1, 0).text = "Controle"
    table.cell(1, 1).text = "12.5"
    table.cell(1, 2).text = "12.5"
    
    table.cell(2, 0).text = "Tratado"
    table.cell(2, 1).text = "18.3"
    table.cell(2, 2).text = "18.3"
    
    # Aplica formatação
    num_formatted = format_all_tables_in_docx(doc)
    print(f"Formatadas {num_formatted} tabelas com sucesso.")
    
    output_path = "test_table_output.docx"
    doc.save(output_path)
    print(f"Arquivo salvo em {output_path}")

if __name__ == "__main__":
    test_table_formatting()
