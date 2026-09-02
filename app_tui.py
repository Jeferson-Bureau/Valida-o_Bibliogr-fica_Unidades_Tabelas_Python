"""
app_tui.py — Interface TUI moderna para o Validador de Referências Bibliográficas
Executar com: python app_tui.py
"""
import os
import csv
import sys
import threading
from typing import List, Tuple
import pyperclip
import tkinter as tk
from tkinter import filedialog

from textual import on, work
from textual.app import App, ComposeResult
from textual.binding import Binding
from textual.containers import Container, Horizontal, Vertical, ScrollableContainer
from textual.reactive import reactive
from textual.widgets import (
    Button,
    Footer,
    Header,
    Input,
    Label,
    Log,
    ProgressBar,
    Static,
    Rule,
)

from models import ReferenceInput, ValidationResult
from word_parser import extract_references_from_docx
from reference_parser import parse_abnt_reference
from main import ReferenceValidator
from apa7_formatter import format_apa7
from si_units_corrector import correct_si_units
from table_formatter import format_all_tables_in_docx

# ─────────────────────────────────────────────────────────────
# CSS da aplicação
# ─────────────────────────────────────────────────────────────
APP_CSS = """
Screen {
    background: $surface;
}

Header {
    background: #1a3a5c;
    color: #e0f0ff;
}

#layout {
    layout: horizontal;
    height: 1fr;
}

/* ── Painel esquerdo ── */
#left-panel {
    width: 36;
    min-width: 32;
    background: #0f2236;
    border-right: solid #1a3a5c;
    padding: 1 2;
}

#app-title {
    text-align: center;
    color: #7ec8e3;
    text-style: bold;
    margin-bottom: 1;
}

#label-file, #label-output {
    color: #7ec8e3;
    text-style: bold;
    margin-top: 1;
    margin-bottom: 0;
}

#file-input {
    width: 100%;
    margin-bottom: 1;
    background: #162d45;
    border: solid #1a3a5c;
    color: #e0f0ff;
}

#output-input {
    width: 100%;
    margin-bottom: 1;
    background: #162d45;
    border: solid #1a3a5c;
    color: #e0f0ff;
}

#file-input:focus, #output-input:focus {
    border: solid #7ec8e3;
}

#btn-browse {
    width: 100%;
    margin-top: 0;
    background: #1a3d6b;
    color: #aad4ff;
    border: solid #2255aa;
}

#btn-browse:hover {
    background: #2255aa;
}

#btn-paste {
    width: 100%;
    margin-top: 0;
    background: #2a2a4a;
    color: #aabbff;
    border: solid #4444aa;
}

#btn-paste:hover {
    background: #3a3a6a;
}

#btn-validate {
    width: 100%;
    margin-top: 1;
    background: #1a6b3c;
    color: #ffffff;
    text-style: bold;
    border: solid #2ea055;
}

#btn-validate:hover {
    background: #2ea055;
}

#btn-csv {
    width: 100%;
    margin-top: 1;
    background: #163d6b;
    color: #e0f0ff;
    border: solid #1a3a5c;
}

#btn-csv:hover {
    background: #1a6b8f;
}

#btn-si {
    width: 100%;
    margin-top: 1;
    background: #3b1a5c;
    color: #e0ccff;
    border: solid #7744bb;
}

#btn-si:hover {
    background: #5c2ea0;
}

#btn-tables {
    width: 100%;
    margin-top: 1;
    background: #1a5c54;
    color: #ccffff;
    border: solid #2ea090;
}

#btn-tables:hover {
    background: #2ea090;
}

#btn-quit {
    width: 100%;
    margin-top: 1;
    background: #5c1a1a;
    color: #ffcccc;
    border: solid #8b2020;
}

#btn-quit:hover {
    background: #8b2020;
}

/* ── Painel direito ── */
#right-panel {
    width: 1fr;
    layout: vertical;
    background: $surface;
}

/* ── Log ── */
#log-container {
    height: 1fr;
    border: solid #1a3a5c;
    margin: 1 1 0 1;
    background: #080f18;
}

#log-label {
    background: #1a3a5c;
    color: #7ec8e3;
    text-style: bold;
    padding: 0 1;
}

#validation-log {
    background: #080f18;
    color: #c0d8f0;
    height: 1fr;
}

/* ── Progresso ── */
#progress-container {
    height: auto;
    border: solid #1a3a5c;
    margin: 1;
    padding: 0 1 1 1;
    background: #0f1e2e;
}

#progress-label {
    color: #7ec8e3;
    text-style: bold;
    margin-bottom: 0;
}

#progress-bar {
    width: 100%;
    margin-top: 0;
}

#progress-info {
    color: #88aacc;
    text-align: center;
}

/* ── Resumo ── */
#summary-container {
    height: auto;
    layout: horizontal;
    margin: 0 1 1 1;
}

.summary-box {
    width: 1fr;
    height: 5;
    border: solid #1a3a5c;
    padding: 0 1;
    content-align: center middle;
    text-align: center;
}

#box-confirmed {
    background: #0d2e18;
    border: solid #2ea055;
}

#box-partial {
    background: #2e1f00;
    border: solid #cc6600;
}

#box-notfound {
    background: #2e0d0d;
    border: solid #8b2020;
}

.summary-count {
    text-style: bold;
    text-align: center;
}

#count-confirmed { color: #2ea055; }
#count-partial   { color: #cc8800; }
#count-notfound  { color: #cc3333; }

.summary-title {
    text-align: center;
    color: #88aacc;
}
"""

# ─────────────────────────────────────────────────────────────
# Widgets auxiliares
# ─────────────────────────────────────────────────────────────
class SummaryBox(Static):
    """Caixinha de resumo (Confirmado / Parcial / Nao confirmado)."""

    def __init__(self, box_id: str, title: str, count_id: str, initial: str = "0", **kwargs):
        super().__init__(**kwargs)
        self._box_id = box_id
        self._title = title
        self._count_id = count_id
        self._initial = initial

    def compose(self) -> ComposeResult:
        yield Label(self._initial, id=self._count_id, classes="summary-count")
        yield Label(self._title, classes="summary-title")


# ─────────────────────────────────────────────────────────────
# Aplicação principal
# ─────────────────────────────────────────────────────────────
class ValidadorApp(App):
    """TUI do Validador de Referencias Bibliograficas."""

    CSS = APP_CSS
    TITLE = "Validador de Referencias Bibliograficas"
    BINDINGS = [
        Binding("q", "quit", "Sair"),
        Binding("f1", "show_help", "Ajuda"),
        Binding("f5", "start_validation", "Validar [F5]"),
    ]

    # Contadores reativos
    confirmed = reactive(0)
    partial = reactive(0)
    not_found = reactive(0)
    total = reactive(0)
    processed = reactive(0)

    # ── Layout ──────────────────────────────────────────────
    def compose(self) -> ComposeResult:
        yield Header(show_clock=True)

        with Horizontal(id="layout"):

            # Painel esquerdo
            with Vertical(id="left-panel"):
                yield Label("Validador Bibliografico", id="app-title")
                yield Rule()
                yield Label("Arquivo Word (.docx):", id="label-file")
                yield Input(
                    placeholder="Use os botoes abaixo para selecionar...",
                    id="file-input",
                )
                yield Button("Procurar arquivo (.docx)...", id="btn-browse", variant="primary")
                yield Button("[ Colar caminho do arquivo ]", id="btn-paste", variant="default")
                yield Label("Salvar resultado como:", id="label-output")
                yield Input(
                    placeholder="(opcional) _verificado.docx",
                    id="output-input",
                )
                yield Rule()
                yield Button(">> VALIDAR REFERENCIAS",    id="btn-validate", variant="success")
                yield Button("   Formatar Tabelas",     id="btn-tables",   variant="default")
                yield Button("   Corrigir Unidades SI", id="btn-si",       variant="default")
                yield Button("   Abrir CSV no Excel",   id="btn-csv",      variant="default")
                yield Button("   Sair",                 id="btn-quit",     variant="error")

            # Painel direito
            with Vertical(id="right-panel"):

                # Log
                with Container(id="log-container"):
                    yield Label(" LOG DE VALIDACAO ", id="log-label")
                    yield Log(id="validation-log", auto_scroll=True)

                # Barra de progresso
                with Vertical(id="progress-container"):
                    yield Label("PROGRESSO", id="progress-label")
                    yield ProgressBar(total=100, show_eta=False, id="progress-bar")
                    yield Label("Aguardando arquivo...", id="progress-info")

                # Resumo
                with Horizontal(id="summary-container"):
                    with Container(classes="summary-box", id="box-confirmed"):
                        yield Label("0", id="count-confirmed", classes="summary-count")
                        yield Label("Confirmado", classes="summary-title")
                    with Container(classes="summary-box", id="box-partial"):
                        yield Label("0", id="count-partial", classes="summary-count")
                        yield Label("Parcial", classes="summary-title")
                    with Container(classes="summary-box", id="box-notfound"):
                        yield Label("0", id="count-notfound", classes="summary-count")
                        yield Label("Nao confirmado", classes="summary-title")

        yield Footer()

    # ── Watchers reativos ────────────────────────────────────
    def watch_confirmed(self, value: int):
        try:
            self.query_one("#count-confirmed", Label).update(str(value))
        except Exception:
            pass

    def watch_partial(self, value: int):
        try:
            self.query_one("#count-partial", Label).update(str(value))
        except Exception:
            pass

    def watch_not_found(self, value: int):
        try:
            self.query_one("#count-notfound", Label).update(str(value))
        except Exception:
            pass

    def watch_processed(self, value: int):
        if self.total > 0:
            pct = int((value / self.total) * 100)
            self.query_one("#progress-bar", ProgressBar).update(progress=pct)
            self.query_one("#progress-info", Label).update(
                f"{value} / {self.total} referencias processadas"
            )

    # ── Botões ───────────────────────────────────────────────
    @on(Button.Pressed, "#btn-browse")
    def on_browse(self) -> None:
        """Abre a janela nativa do Windows para selecionar o arquivo .docx."""
        def open_dialog():
            root = tk.Tk()
            root.withdraw()          # Esconde a janela principal do Tk
            root.attributes("-topmost", True)  # Garante que aparece na frente
            filepath = filedialog.askopenfilename(
                title="Selecione o arquivo Word",
                filetypes=[("Documentos Word", "*.docx"), ("Todos os arquivos", "*.*")]
            )
            root.destroy()
            if filepath:
                # Atualiza o campo na thread principal do Textual
                self.call_from_thread(
                    self._set_filepath, filepath.replace("/", "\\")
                )

        # Roda em thread separada para nao travar a TUI
        threading.Thread(target=open_dialog, daemon=True).start()

    def _set_filepath(self, filepath: str) -> None:
        """Atualiza o campo de arquivo e exibe log (chamado da thread principal)."""
        self.query_one("#file-input", Input).value = filepath
        self._log(f"Arquivo selecionado: {filepath}", "info")

    @on(Button.Pressed, "#btn-paste")
    def on_paste_clipboard(self) -> None:
        """Le o caminho da area de transferencia e preenche o campo automaticamente."""
        try:
            text = pyperclip.paste().strip().strip('"')
            if text:
                self.query_one("#file-input", Input).value = text
                self._log(f"Caminho colado: {text}", "info")
            else:
                self._log("[AVISO] Area de transferencia vazia. Copie o caminho do arquivo primeiro.", "warning")
        except Exception as e:
            self._log(f"[ERRO] Nao foi possivel ler o clipboard: {e}", "error")

    @on(Button.Pressed, "#btn-validate")
    def on_validate(self) -> None:
        filepath = self.query_one("#file-input", Input).value.strip().strip('"')
        if not filepath:
            self._log("[ERRO] Informe o caminho do arquivo .docx no campo acima.", "error")
            return
        if not os.path.exists(filepath):
            self._log(f"[ERRO] Arquivo nao encontrado: {filepath}", "error")
            return
        self._start_validation(filepath)

    @on(Button.Pressed, "#btn-tables")
    def on_format_tables(self) -> None:
        """Abre o arquivo .docx e aplica o formatador de tabelas da Acta Scientiarum."""
        filepath = self.query_one("#file-input", Input).value.strip().strip('"')
        if not filepath:
            self._log("[ERRO] Informe o arquivo .docx no campo acima antes de formatar tabelas.", "error")
            return
        if not os.path.exists(filepath):
            self._log(f"[ERRO] Arquivo nao encontrado: {filepath}", "error")
            return
        self._run_table_formatting(filepath)

    @work(thread=True)
    def _run_table_formatting(self, filepath: str) -> None:
        """Aplica a formatacao de tabelas em thread separada."""
        log = self._log
        log("─" * 60, "info")
        log(f"[TABELAS] Iniciando formatacao de tabelas: {os.path.basename(filepath)}", "info")
        log("[TABELAS] Padrão: Acta Scientiarum / APA (Estilo 3 Linhas - 0,5 pt)", "info")
        log("[TABELAS] Fonte: PT Serif 8 pt, Largura: 100%, Alinhamento: Centro", "info")
        log("─" * 60, "info")

        import docx as docx_lib
        try:
            doc = docx_lib.Document(filepath)
        except Exception as e:
            log(f"[ERRO] Nao foi possivel abrir o arquivo: {e}", "error")
            return

        if len(doc.tables) == 0:
            log("[AVISO] Nenhuma tabela encontrada neste documento.", "warning")
            return

        try:
            n = format_all_tables_in_docx(doc)
        except Exception as e:
            log(f"[ERRO] Falha durante formatacao das tabelas: {e}", "error")
            return

        base, ext = os.path.splitext(filepath)
        output_path = f"{base}_Tabelas{ext}"
        output_custom = self.query_one("#output-input", Input).value.strip().strip('"')
        if output_custom:
            output_path = output_custom

        try:
            doc.save(output_path)
            log("─" * 60, "info")
            log(f"[TABELAS] Concluido! {n} tabela(s) formatada(s) com sucesso.", "success")
            log(f"[TABELAS] Arquivo salvo em:", "success")
            log(f"  {output_path}", "success")
            log("─" * 60, "info")
        except Exception as e:
            log(f"[ERRO] Nao foi possivel salvar o arquivo: {e}", "error")

    @on(Button.Pressed, "#btn-si")
    def on_si_units(self) -> None:
        """Abre dialogo para selecionar .docx e aplica Corretor de Unidades SI."""
        filepath = self.query_one("#file-input", Input).value.strip().strip('"')
        if not filepath:
            self._log("[ERRO] Informe o arquivo .docx no campo acima antes de corrigir.", "error")
            return
        if not os.path.exists(filepath):
            self._log(f"[ERRO] Arquivo nao encontrado: {filepath}", "error")
            return
        self._run_si_correction(filepath)

    @work(thread=True)
    def _run_si_correction(self, filepath: str) -> None:
        """Aplica as correcoes de unidades SI em thread separada."""
        log = self._log
        log("─" * 60, "info")
        log(f"[SI] Iniciando correcao de unidades: {os.path.basename(filepath)}", "info")
        log("[SI] Normas: APA 7a Ed. + Sistema Internacional de Unidades (SI/BIPM)", "info")
        log("─" * 60, "info")

        import docx as docx_lib
        try:
            doc = docx_lib.Document(filepath)
        except Exception as e:
            log(f"[ERRO] Nao foi possivel abrir o arquivo: {e}", "error")
            return

        log("[SI] Aplicando correcoes...", "info")
        log("  * Espacamento: numero + unidade (25 kg, 10 mL)", "info")
        log("  * Temperatura: 25 grC  /  Kelvin: 300 K", "info")
        log("  * Porcentagem: 45 % (com espaco, norma APA 7)", "info")
        log("  * Angulos: 30graus (sem espaco)", "info")
        log("  * Unidades compostas: mg/L -> mg L-1", "info")
        log("  * Plurais, prefixos e operadores normalizados", "info")
        log("  * Secao Referencias protegida", "info")

        try:
            n = correct_si_units(doc)
        except Exception as e:
            log(f"[ERRO] Falha durante correcao: {e}", "error")
            return

        # Salvar arquivo corrigido
        base, ext = os.path.splitext(filepath)
        output_path = f"{base}_SI{ext}"
        output_custom = self.query_one("#output-input", Input).value.strip().strip('"')
        if output_custom:
            output_path = output_custom

        try:
            doc.save(output_path)
            log("─" * 60, "info")
            log(f"[SI] Concluido! {n} paragrafos processados.", "success")
            log(f"[SI] Arquivo salvo em:", "success")
            log(f"  {output_path}", "success")
            log("─" * 60, "info")
        except Exception as e:
            log(f"[ERRO] Nao foi possivel salvar o arquivo: {e}", "error")

    @on(Button.Pressed, "#btn-csv")
    def on_open_csv(self) -> None:
        csv_path = os.path.join(
            os.path.dirname(os.path.abspath(__file__)),
            "relatorio_referencias.csv"
        )
        if os.path.exists(csv_path):
            os.startfile(csv_path)
            self._log(f"Abrindo CSV: {csv_path}", "info")
        else:
            self._log("[ERRO] relatorio_referencias.csv nao encontrado. Execute a validacao primeiro.", "error")

    @on(Button.Pressed, "#btn-quit")
    def on_quit(self) -> None:
        self.exit()

    # ── Atalho de teclado ────────────────────────────────────
    def action_start_validation(self) -> None:
        self.on_validate()

    def action_show_help(self) -> None:
        self._log("─" * 60, "info")
        self._log("AJUDA:", "info")
        self._log("  1. Cole o caminho do arquivo .docx no campo 'Arquivo Word'.", "info")
        self._log("  2. Clique em 'VALIDAR REFERENCIAS' ou pressione Ctrl+V.", "info")
        self._log("  3. Acompanhe o progresso no painel de log.", "info")
        self._log("  4. Ao terminar, clique em 'Abrir CSV no Excel'.", "info")
        self._log("  5. O Word anotado e salvo com sufixo _verificado.docx.", "info")
        self._log("─" * 60, "info")

    # ── Lógica de validação (thread separada) ─────────────────
    @work(thread=True)
    def _start_validation(self, filepath: str) -> None:
        log = self._log
        log("─" * 60, "info")
        log(f"Iniciando: {os.path.basename(filepath)}", "info")
        log("─" * 60, "info")

        # Resetar contadores
        self.confirmed = 0
        self.partial = 0
        self.not_found = 0
        self.processed = 0
        self.total = 0

        try:
            doc, ref_paragraphs = extract_references_from_docx(filepath)
        except Exception as e:
            log(f"[ERRO] Nao foi possivel abrir o arquivo: {e}", "error")
            return

        if not ref_paragraphs:
            log("[AVISO] Nenhuma referencia encontrada.", "warning")
            log("  Verifique se a secao se chama exatamente 'Referencias'.", "warning")
            return

        self.total = len(ref_paragraphs)
        log(f"{self.total} referencia(s) encontrada(s). Consultando APIs...", "info")

        validator = ReferenceValidator()
        results: List[Tuple[ReferenceInput, ValidationResult]] = []

        for para_obj, ref_text in ref_paragraphs:
            short = ref_text[:65] + "..." if len(ref_text) > 65 else ref_text
            log(f"  >> {short}", "info")

            parsed = parse_abnt_reference(ref_text)
            ref = ReferenceInput(
                id=ref_text[:40],
                original_text=ref_text,
                title=parsed.get("title"),
                doi=parsed.get("doi"),
                authors=parsed.get("authors", [])
            )

            res = validator.validate_reference(ref)
            results.append((ref, res))

            # Atualizar contadores e log por status
            if res.status == "Confirmado":
                self.confirmed += 1
                log(f"     [OK] Confirmado via {res.api_source}", "success")
                if res.doi_found:
                    log(f"          DOI: {res.doi_found}", "info")
                apa7 = format_apa7(res)
                if apa7:
                    log(f"          APA7: {apa7}", "info")
            elif "Parcial" in res.status:
                self.partial += 1
                log(f"     [!!] Parcialmente confirmado via {res.api_source}", "warning")
                for issue in res.issues:
                    log(f"          - {issue}", "warning")
                apa7 = format_apa7(res)
                if apa7:
                    log(f"          APA7: {apa7}", "info")
            else:
                self.not_found += 1
                log(f"     [XX] Nao confirmado em nenhuma base", "error")

            # Anotar diretamente no parágrafo correto da lista de referências
            if doc and para_obj is not None:
                from main import annotate_paragraph
                annotate_paragraph(para_obj, res)

            self.processed += 1

        # Salvar Word anotado
        output_path = self.query_one("#output-input", Input).value.strip().strip('"')
        if not output_path:
            base, ext = os.path.splitext(filepath)
            output_path = f"{base}_verificado{ext}"

        if doc:
            try:
                doc.save(output_path)
                log("─" * 60, "info")
                log(f"Word verificado salvo em:", "success")
                log(f"  {output_path}", "success")
            except Exception as e:
                log(f"[ERRO] Nao foi possivel salvar o Word: {e}", "error")

        # Salvar CSV
        csv_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "relatorio_referencias.csv")
        try:
            with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
                writer = csv.writer(f)
                writer.writerow(["N", "Referencia Original", "Status", "Autoria",
                                  "DOI/Link", "Autores Encontrados", "Fonte", "Confianca",
                                  "Observacoes", "Referencia APA7"])
                for idx, (ref, res) in enumerate(results, 1):
                    autoria = "Sim" if res.status == "Confirmado" else ("Parcial" if "Parcial" in res.status else "Nao")
                    apa7 = format_apa7(res) or "-"
                    writer.writerow([
                        idx, ref.original_text, res.status, autoria,
                        res.doi_found or res.url or "-",
                        "; ".join(res.authors_found),
                        res.api_source, res.confidence,
                        " | ".join(res.issues) if res.issues else "-",
                        apa7,
                    ])
            log(f"Relatorio CSV salvo em:", "success")
            log(f"  {csv_path}", "success")
        except Exception as e:
            log(f"[ERRO] Nao foi possivel salvar o CSV: {e}", "error")

        log("─" * 60, "info")
        log(f"CONCLUIDO: {self.confirmed} confirmadas / {self.partial} parciais / {self.not_found} nao confirmadas", "success")

    # ── Helper de log colorido ────────────────────────────────
    def _log(self, message: str, level: str = "info") -> None:
        widget = self.query_one("#validation-log", Log)
        prefix_map = {
            "info":    "",
            "success": "[OK] ",
            "warning": "[!!] ",
            "error":   "[XX] ",
        }
        prefix = prefix_map.get(level, "")
        widget.write_line(f"{prefix}{message}")


# ─────────────────────────────────────────────────────────────
# Entrada
# ─────────────────────────────────────────────────────────────
if __name__ == "__main__":
    app = ValidadorApp()
    app.run()
