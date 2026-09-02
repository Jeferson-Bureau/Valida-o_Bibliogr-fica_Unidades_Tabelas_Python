"""
table_style.py
================
Descreve a formatação de uma tabela do Word como dados estruturados
(dataclasses / JSON) e aplica essa formatação a tabelas de um .docx
usando python-docx (com acesso ao XML de baixo nível para bordas,
sombreamento e margens de célula, que a API pública não expõe).

Uso básico:

    from docx import Document
    from table_style import TableStyleSpec, apply_table_style

    doc = Document("entrada.docx")
    estilo = TableStyleSpec()          # preset "três linhas" (Acta Scientiarum)
    for tabela in doc.tables:
        apply_table_style(tabela, estilo)
    doc.save("saida.docx")

O objeto TableStyleSpec pode ser serializado/lido como JSON, o que
permite guardar "perfis" de formatação (um por revista, por template
etc.) sem tocar em código.
"""

from dataclasses import dataclass, field, asdict
from typing import Optional, List
import json

from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL


# ======================================================================
# 1) SCHEMA — descreve a formatação de forma independente do python-docx
# ======================================================================

@dataclass
class BorderSpec:
    """Uma borda simples no padrão OOXML."""
    style: str = "single"        # "single" | "double" | "dashed" | ...
    size_eighths_pt: int = 4     # unidade OOXML = 1/8 pt -> 4 = 0,5 pt
    color: str = "auto"          # "auto" ou hex "RRGGBB"


@dataclass
class TableBordersSpec:
    """
    Bordas por posição lógica na tabela (não por linha/coluna física),
    para funcionar em tabelas de qualquer tamanho.
    """
    top: Optional[BorderSpec] = field(default_factory=BorderSpec)            # acima da 1ª linha
    header_bottom: Optional[BorderSpec] = field(default_factory=BorderSpec)  # sob o cabeçalho
    bottom: Optional[BorderSpec] = field(default_factory=BorderSpec)         # abaixo da última linha
    left: Optional[BorderSpec] = None
    right: Optional[BorderSpec] = None
    inside_horizontal: Optional[BorderSpec] = None   # entre linhas de dados
    inside_vertical: Optional[BorderSpec] = None      # entre colunas


@dataclass
class FontSpec:
    name: str = "PT Serif"
    size_pt: float = 8.0
    bold: bool = False
    color_hex: str = "000000"


@dataclass
class TableStyleSpec:
    """Especificação completa de formatação de uma tabela."""
    width_pct: float = 100.0                       # largura da tabela (% da área de texto)
    column_widths_pct: Optional[List[float]] = None  # opcional: largura relativa de cada coluna
    cell_margin_top_pt: float = 0.0                 # Margem superior 0 cm
    cell_margin_bottom_pt: float = 0.0              # Margem inferior 0 cm
    cell_margin_left_pt: float = 0.0                # Margem esquerda 0 cm (conforme caixa de diálogo "Opções de Célula")
    cell_margin_right_pt: float = 0.0               # Margem direita 0 cm (conforme caixa de diálogo "Opções de Célula")
    wrap_text: bool = True                          # Quebrar texto automaticamente (w:noWrap ausente/desativado)
    fit_text: bool = False                          # Ajustar texto (w:tcFitText)
    borders: TableBordersSpec = field(default_factory=TableBordersSpec)
    font: FontSpec = field(default_factory=FontSpec)
    align_horizontal: str = "center"                # "left" | "center" | "right"
    align_vertical: str = "center"                  # "top" | "center" | "bottom"
    header_rows: int = 1                            # nº de linhas tratadas como cabeçalho
    remove_shading: bool = True                      # remove qualquer sombreamento de fundo
    dont_add_space_between_same_style: bool = False # "Não adicionar espaço entre parágrafos do mesmo estilo" (desativado na caixa)
    line_spacing_type: str = "single"                # "single" (Espaçamento entre linhas: Simples)
    space_before_pt: float = 0.0                     # Espaçamento Antes: 0 pt
    space_after_pt: float = 0.0                      # Espaçamento Depois: 0 pt
    indent_left_pt: float = 0.0                      # Recuo Esquerda: 0 cm
    indent_right_pt: float = 0.0                     # Recuo Direita: 0 cm
    widow_control: bool = True                       # Controle de linhas órfãs/viúvas (Ativado)
    keep_with_next: bool = False                    # Manter com o próximo (Desativado)
    keep_lines: bool = False                        # Manter linhas juntas (Desativado)
    page_break_before: bool = False                 # Quebrar página antes (Desativado)
    suppress_line_numbers: bool = False             # Suprimir números de linha (Desativado)
    suppress_auto_hyphens: bool = True              # Não hifenizar (Ativado -> w:suppressAutoHyphens)
    min_row_height_cm: Optional[float] = 0.00        # Altura mínima da linha em cm (0,00 cm)
    distribute_columns: bool = True                  # Distribuir colunas igualmente


    # ---- (de)serialização para guardar "perfis" de formatação ----
    def to_json(self) -> str:
        return json.dumps(asdict(self), indent=2, ensure_ascii=False)

    @classmethod
    def from_json(cls, text: str) -> "TableStyleSpec":
        data = json.loads(text)
        borders_data = data.pop("borders")
        font_data = data.pop("font")
        borders = TableBordersSpec(**{
            k: (BorderSpec(**v) if v is not None else None)
            for k, v in borders_data.items()
        })
        font = FontSpec(**font_data)
        return cls(borders=borders, font=font, **data)


# Preset equivalente à tabela "Table 1. Ice cream formulations."
# (Acta Scientiarum / padrão de "três linhas")
ACTA_SCIENTIARUM_STYLE = TableStyleSpec()


# ======================================================================
# 2) APLICAÇÃO — funções auxiliares de baixo nível (OOXML) + função
#    pública apply_table_style()
# ======================================================================

def _get_or_add(parent, tag: str):
    el = parent.find(qn(tag))
    if el is None:
        el = OxmlElement(tag)
        parent.append(el)
    return el


def _set_cell_border(cell, edges: dict):
    """edges: {'top'|'bottom'|'left'|'right': BorderSpec | None}"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = _get_or_add(tcPr, "w:tcBorders")
    for edge, spec in edges.items():
        el = _get_or_add(tcBorders, f"w:{edge}")
        if spec is None:
            el.set(qn("w:val"), "nil")
            for attr in ("w:sz", "w:space", "w:color"):
                if el.get(qn(attr)) is not None:
                    del el.attrib[qn(attr)]
        else:
            el.set(qn("w:val"), spec.style)
            el.set(qn("w:sz"), str(spec.size_eighths_pt))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), spec.color)


def _set_cell_margins(cell, top_pt, bottom_pt, left_pt, right_pt):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = _get_or_add(tcPr, "w:tcMar")
    for edge, val_pt in (("top", top_pt), ("bottom", bottom_pt),
                          ("left", left_pt), ("right", right_pt)):
        el = _get_or_add(tcMar, f"w:{edge}")
        el.set(qn("w:w"), str(int(round(val_pt * 20))))  # pt -> dxa (1pt = 20dxa)
        el.set(qn("w:type"), "dxa")


def _clear_cell_shading(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = _get_or_add(tcPr, "w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "auto")


def _set_table_width_pct(table, pct: float):
    tblPr = table._tbl.tblPr
    tblW = _get_or_add(tblPr, "w:tblW")
    tblW.set(qn("w:type"), "pct")
    tblW.set(qn("w:w"), str(int(pct * 50)))  # OOXML: 5000 = 100%


def _set_column_widths_pct(table, widths_pct: List[float]):
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    cols = grid.findall(qn("w:gridCol"))
    # largura total atual da grade (dxa), usada como referência proporcional
    total_dxa = sum(int(c.get(qn("w:w")) or 0) for c in cols) or 9000
    for col, pct in zip(cols, widths_pct):
        col.set(qn("w:w"), str(int(total_dxa * pct / 100)))


_VALIGN = {
    "top": WD_ALIGN_VERTICAL.TOP,
    "center": WD_ALIGN_VERTICAL.CENTER,
    "bottom": WD_ALIGN_VERTICAL.BOTTOM,
}
_HALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
}


def apply_table_style(table, spec: TableStyleSpec) -> None:
    """Aplica `spec` a uma única tabela (docx.table.Table)."""
    n_rows = len(table.rows)

    _set_table_width_pct(table, spec.width_pct)
    if spec.column_widths_pct:
        _set_column_widths_pct(table, spec.column_widths_pct)
    elif spec.distribute_columns:
        n_cols = len(table.columns) if hasattr(table, 'columns') and len(table.columns) > 0 else (len(table.rows[0].cells) if n_rows > 0 else 0)
        if n_cols > 0:
            equal_pct = [100.0 / n_cols] * n_cols
            _set_column_widths_pct(table, equal_pct)

    for r_idx, row in enumerate(table.rows):
        is_header = r_idx < spec.header_rows
        is_last = r_idx == n_rows - 1

        top_border = spec.borders.top if r_idx == 0 else spec.borders.inside_horizontal
        if is_header:
            bottom_border = spec.borders.header_bottom
        elif is_last:
            bottom_border = spec.borders.bottom
        else:
            bottom_border = spec.borders.inside_horizontal

        if spec.min_row_height_cm is not None:
            # 1 cm = 567 dxa aproximadamente (28.35 pt * 20 = 567 dxa)
            height_dxa = int(round(spec.min_row_height_cm * 567))
            trPr = _get_or_add(row._tr, "w:trPr")
            trHeight = _get_or_add(trPr, "w:trHeight")
            trHeight.set(qn("w:val"), str(height_dxa))
            trHeight.set(qn("w:hRule"), "atLeast")
        elif spec.row_height_auto:
            trPr = _get_or_add(row._tr, "w:trPr")
            trHeight = _get_or_add(trPr, "w:trHeight")
            trHeight.attrib.pop(qn("w:hRule"), None)

        for cell in row.cells:
            _set_cell_border(cell, {
                "top": top_border,
                "bottom": bottom_border,
                "left": spec.borders.left,
                "right": spec.borders.right,
            })
            _set_cell_margins(
                cell,
                spec.cell_margin_top_pt, spec.cell_margin_bottom_pt,
                spec.cell_margin_left_pt, spec.cell_margin_right_pt,
            )
            if spec.remove_shading:
                _clear_cell_shading(cell)

            tcPr = cell._tc.get_or_add_tcPr()
            if spec.wrap_text:
                noWrap = tcPr.find(qn("w:noWrap"))
                if noWrap is not None:
                    tcPr.remove(noWrap)
            else:
                _get_or_add(tcPr, "w:noWrap")

            if spec.fit_text:
                _get_or_add(tcPr, "w:tcFitText")
            else:
                tcFitText = tcPr.find(qn("w:tcFitText"))
                if tcFitText is not None:
                    tcPr.remove(tcFitText)

            cell.vertical_alignment = _VALIGN[spec.align_vertical]

            for p in cell.paragraphs:
                p.alignment = _HALIGN[spec.align_horizontal]
                p.paragraph_format.space_before = Pt(spec.space_before_pt)
                p.paragraph_format.space_after = Pt(spec.space_after_pt)
                p.paragraph_format.left_indent = Pt(spec.indent_left_pt)
                p.paragraph_format.right_indent = Pt(spec.indent_right_pt)
                p.paragraph_format.first_line_indent = Pt(0)
                
                # Espaçamento entre linhas: Simples (1.0)
                if spec.line_spacing_type == "single":
                    p.paragraph_format.line_spacing = 1.0

                # "Não adicionar espaço entre parágrafos do mesmo estilo" -> w:contextualSpacing
                pPr = p._p.get_or_add_pPr()
                ctxSp = pPr.find(qn("w:contextualSpacing"))
                if spec.dont_add_space_between_same_style:
                    if ctxSp is None:
                        _get_or_add(pPr, "w:contextualSpacing")
                else:
                    if ctxSp is not None:
                        pPr.remove(ctxSp)

                # --- Quebras de linha e de página / Paginação ---
                # 1. Controle de linhas órfãs/viúvas
                wc = pPr.find(qn("w:widowControl"))
                if spec.widow_control:
                    if wc is None:
                        _get_or_add(pPr, "w:widowControl")
                else:
                    if wc is not None:
                        pPr.remove(wc)

                # 2. Manter com o próximo
                kn = pPr.find(qn("w:keepNext"))
                if spec.keep_with_next:
                    if kn is None:
                        _get_or_add(pPr, "w:keepNext")
                else:
                    if kn is not None:
                        pPr.remove(kn)

                # 3. Manter linhas juntas
                kl = pPr.find(qn("w:keepLines"))
                if spec.keep_lines:
                    if kl is None:
                        _get_or_add(pPr, "w:keepLines")
                else:
                    if kl is not None:
                        pPr.remove(kl)

                # 4. Quebrar página antes
                pb = pPr.find(qn("w:pageBreakBefore"))
                if spec.page_break_before:
                    if pb is None:
                        _get_or_add(pPr, "w:pageBreakBefore")
                else:
                    if pb is not None:
                        pPr.remove(pb)

                # --- Exceções de formatação ---
                # 5. Suprimir números de linha
                sln = pPr.find(qn("w:suppressLineNumbers"))
                if spec.suppress_line_numbers:
                    if sln is None:
                        _get_or_add(pPr, "w:suppressLineNumbers")
                else:
                    if sln is not None:
                        pPr.remove(sln)

                # 6. Não hifenizar (w:suppressAutoHyphens)
                sah = pPr.find(qn("w:suppressAutoHyphens"))
                if spec.suppress_auto_hyphens:
                    if sah is None:
                        _get_or_add(pPr, "w:suppressAutoHyphens")
                else:
                    if sah is not None:
                        pPr.remove(sah)

                if not p.runs and p.text == "":
                    continue
                for run in p.runs:
                    run.font.name = spec.font.name
                    run.font.size = Pt(spec.font.size_pt)
                    run.font.bold = spec.font.bold
                    run.font.color.rgb = RGBColor.from_string(spec.font.color_hex)


def apply_to_all_tables(doc, spec: TableStyleSpec) -> int:
    """Aplica `spec` a todas as tabelas de um docx.Document. Retorna a contagem."""
    count = 0
    for table in doc.tables:
        apply_table_style(table, spec)
        count += 1
    return count


# ======================================================================
# 3) Exemplo de uso via linha de comando
# ======================================================================
if __name__ == "__main__":
    import argparse
    from docx import Document

    parser = argparse.ArgumentParser(
        description="Aplica formatação padronizada a todas as tabelas de um .docx"
    )
    parser.add_argument("entrada", help="Arquivo .docx de entrada")
    parser.add_argument("saida", help="Arquivo .docx de saída")
    parser.add_argument("--perfil", help="Arquivo JSON com um TableStyleSpec customizado")
    args = parser.parse_args()

    estilo = (
        TableStyleSpec.from_json(open(args.perfil, encoding="utf-8").read())
        if args.perfil else ACTA_SCIENTIARUM_STYLE
    )

    documento = Document(args.entrada)
    total = apply_to_all_tables(documento, estilo)
    documento.save(args.saida)
    print(f"{total} tabela(s) formatada(s). Arquivo salvo em: {args.saida}")
